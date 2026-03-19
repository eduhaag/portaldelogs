"""
===============================
Portal do Suporte - Backend Principal
===============================
Aqui é onde a mágica acontece! Este arquivo é o maestro que orquestra APIs, autenticação, conexão com banco, análise de logs e até exportação de relatórios bonitos.
Prepare-se para encontrar endpoints, modelos, integrações e um toque de bom humor nos comentários!
"""

import asyncio
import fastapi
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse, Response
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
import hashlib
from pathlib import Path
from pydantic import BaseModel, Field
from typing import Any, Dict, List, Optional
import uuid
from datetime import datetime, timedelta
import io
import re
import json
import zipfile
import csv
import html
import tempfile
import shutil
from time import monotonic
from jose import JWTError, jwt
from log_analyzer import LogAnalyzer
from large_log_processor import LargeLogProcessor, process_large_log_file
from local_pattern_store import delete_records, find_record, insert_record, list_records, update_records
from version_compare_service import version_compare_service
import aiofiles
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, BaseDocTemplate, PageTemplate, Frame
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from xml.sax.saxutils import escape
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from lxml import html as lxml_html


ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')


# ===============================
# MongoDB: onde os logs vão para serem analisados e julgados!
# ===============================
mongo_url = os.environ['MONGO_URL']
mongo_server_selection_timeout_ms = int(os.environ.get('MONGO_SERVER_SELECTION_TIMEOUT_MS', '1500'))
mongo_connect_timeout_ms = int(os.environ.get('MONGO_CONNECT_TIMEOUT_MS', '1500'))
mongo_socket_timeout_ms = int(os.environ.get('MONGO_SOCKET_TIMEOUT_MS', '3000'))
client = AsyncIOMotorClient(
    mongo_url,
    serverSelectionTimeoutMS=mongo_server_selection_timeout_ms,
    connectTimeoutMS=mongo_connect_timeout_ms,
    socketTimeoutMS=mongo_socket_timeout_ms,
)
db = client[os.environ['DB_NAME']]

# ===============================
# Segurança: tokens JWT para proteger seu precioso backend
# ===============================
JWT_SECRET_KEY = os.environ.get('JWT_SECRET_KEY', 'central-suporte-dev-secret-change-me')
JWT_ALGORITHM = 'HS256'
JWT_EXPIRE_MINUTES = int(os.environ.get('JWT_ACCESS_TOKEN_EXPIRE_MINUTES', '480'))
PUBLIC_API_PATHS = {
    '/api/',
    '/api/auth/login',
    '/api/auth/register',
    '/api/status',
}


# ===============================
# FastAPI: o chefão das rotas!
# ===============================
app = fastapi.FastAPI()

# Roteador com prefixo /api para organizar a bagunça
api_router = fastapi.APIRouter(prefix="/api")



# ===============================
# Modelos Pydantic: porque até os dados precisam de regras!
# ===============================
class StatusCheck(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_name: str
    timestamp: datetime = Field(default_factory=datetime.utcnow)

class StatusCheckCreate(BaseModel):
    client_name: str

class LogAnalysisResult(BaseModel):
    success: bool
    log_type: Optional[str] = None
    statistics: dict
    results: List[dict]
    total_results: int
    chart_data: dict
    error_counts: dict
    severity_counts: dict
    attention_points: Optional[List[dict]] = None
    total_attention_points: Optional[int] = None
    informational_lines: Optional[List[dict]] = None
    new_errors: Optional[dict] = None
    performance_analysis: Optional[dict] = None
    top_programs_methods: Optional[dict] = None
    structured_analysis: Optional[dict] = None  # Dados do parsing estruturado
    analysis_timings: Optional[dict] = None
    error: Optional[str] = None

class AddPatternRequest(BaseModel):
    pattern: str
    partial_pattern: Optional[str] = None
    description: str
    category: str
    severity: str
    example: Optional[str] = None
    solution: str

class SearchResult(BaseModel):
    success: bool
    total_matches: int
    matches: List[dict]
    search_info: dict
    error: Optional[str] = None

class ErrorCategorization(BaseModel):
    pattern_id: str
    pattern: str
    category_type: str  # "permanent" or "session"
    description: Optional[str] = None

class NonErrorPattern(BaseModel):
    pattern: str
    full_message: str
    partial_pattern: Optional[str] = None
    reason: Optional[str] = None
    source_line: Optional[int | str] = None


class AuthUserCreate(BaseModel):
    display_name: str
    username: str
    email: str
    password: str


class AuthLoginRequest(BaseModel):
    username: str
    password: str


class AuthUserResponse(BaseModel):
    username: str
    display_name: str
    email: str


class AuthResult(BaseModel):
    success: bool
    message: str
    user: Optional[AuthUserResponse] = None
    access_token: Optional[str] = None
    token_type: Optional[str] = None
    expires_in: Optional[int] = None
    expires_at: Optional[str] = None


def decode_uploaded_text(content: bytes) -> str:
    """Converte bytes enviados em upload para texto, com fallback de encoding."""
    try:
        return content.decode('utf-8')
    except UnicodeDecodeError:
        try:
            return content.decode('latin-1')
        except UnicodeDecodeError:
            return content.decode('utf-8', errors='ignore')

# Add your routes to the router instead of directly to app
@api_router.get("/")
async def root():
    return {"message": "Hello World"}

@api_router.post("/status", response_model=StatusCheck)
async def create_status_check(input: StatusCheckCreate):
    status_dict = input.dict()
    status_obj = StatusCheck(**status_dict)
    await save_status_check_record(status_obj.dict())
    return status_obj

@api_router.get("/status", response_model=List[StatusCheck])
async def get_status_checks():
    status_checks = await load_status_check_records(1000)
    return [StatusCheck(**status_check) for status_check in status_checks]

# Log Analysis Endpoints
analyzer = LogAnalyzer()
force_local_store = False
last_auth_store_probe_at = 0.0
AUTH_STORE_PROBE_TIMEOUT_SECONDS = float(os.environ.get('AUTH_STORE_PROBE_TIMEOUT_SECONDS', '0.2'))
AUTH_STORE_PROBE_INTERVAL_SECONDS = float(os.environ.get('AUTH_STORE_PROBE_INTERVAL_SECONDS', '15'))


def _is_local_store_error(error: Exception) -> bool:
    message = str(error).lower()
    return any(
        text in message
        for text in (
            "requires authentication",
            "unauthorized",
            "authentication failed",
            "server selection timeout",
            "connection refused",
        )
    )


def _should_use_local_store(error: Exception) -> bool:
    global force_local_store

    if force_local_store:
        return True

    if not _is_local_store_error(error):
        return False

    force_local_store = True
    logger.warning("MongoDB indisponível para escrita/leitura. Alternando para armazenamento local temporariamente.")
    return True


async def _ensure_auth_store_ready() -> bool:
    """Executa um probe rápido antes do login/cadastro para evitar a primeira espera longa."""
    global last_auth_store_probe_at

    if force_local_store:
        return True

    now = monotonic()
    if last_auth_store_probe_at and (now - last_auth_store_probe_at) < AUTH_STORE_PROBE_INTERVAL_SECONDS:
        return False

    last_auth_store_probe_at = now

    try:
        await asyncio.wait_for(client.admin.command('ping'), timeout=AUTH_STORE_PROBE_TIMEOUT_SECONDS)
        return False
    except asyncio.TimeoutError:
        timeout_error = TimeoutError('server selection timeout during auth probe')
        _should_use_local_store(timeout_error)
        logger.warning('MongoDB não respondeu ao probe rápido de autenticação. Usando armazenamento local.')
        return True
    except Exception as e:
        if _should_use_local_store(e):
            logger.warning(f'MongoDB indisponível no probe rápido de autenticação: {e}')
            return True
        return False


async def load_active_custom_patterns(limit: int = 1000) -> List[Dict[str, Any]]:
    try:
        return await db.custom_patterns.find({"active": True}).sort("created_at", -1).to_list(limit)
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Using local fallback for custom patterns: {e}")
        return list_records("custom_patterns", {"active": True}, sort_field="created_at", descending=True, limit=limit)


async def load_active_non_error_patterns(limit: int = 1000) -> List[Dict[str, Any]]:
    try:
        return await db.non_error_patterns.find({"active": True}).sort("created_at", -1).to_list(limit)
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Using local fallback for non-error patterns: {e}")
        return list_records("non_error_patterns", {"active": True}, sort_field="created_at", descending=True, limit=limit)


async def save_custom_pattern_record(record: Dict[str, Any]) -> Dict[str, Any]:
    try:
        await db.custom_patterns.insert_one(record)
        return record
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Saving custom pattern to local fallback store: {e}")
        return insert_record("custom_patterns", record)


async def save_session_pattern_record(record: Dict[str, Any]) -> Dict[str, Any]:
    try:
        await db.session_patterns.insert_one(record)
        return record
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Saving session pattern to local fallback store: {e}")
        return insert_record("session_patterns", record)


async def save_non_error_pattern_record(record: Dict[str, Any]) -> Dict[str, Any]:
    try:
        await db.non_error_patterns.insert_one(record)
        return record
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Saving non-error pattern to local fallback store: {e}")
        return insert_record("non_error_patterns", record)


async def deactivate_custom_pattern_record(pattern_id: str) -> int:
    try:
        result = await db.custom_patterns.update_one(
            {"id": pattern_id},
            {"$set": {"active": False}}
        )
        return result.matched_count
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Deactivating custom pattern in local fallback store: {e}")
        return update_records("custom_patterns", {"id": pattern_id}, {"active": False})


async def search_custom_patterns_records(search_term: str, limit: int = 10) -> List[Dict[str, Any]]:
    search_terms = [term.replace("*", "").strip().lower() for term in search_term.split() if len(term.strip()) > 2]

    try:
        flexible_regex = "|".join(search_terms)
        custom_search_query = {
            "$or": [
                {"pattern": {"$regex": search_term, "$options": "i"}},
                {"description": {"$regex": search_term, "$options": "i"}},
                {"solution": {"$regex": search_term, "$options": "i"}},
                {"category": {"$regex": search_term, "$options": "i"}},
                {"example": {"$regex": search_term, "$options": "i"}},
                {"pattern": {"$regex": flexible_regex, "$options": "i"}},
                {"description": {"$regex": flexible_regex, "$options": "i"}},
                {"solution": {"$regex": flexible_regex, "$options": "i"}},
                {"example": {"$regex": flexible_regex, "$options": "i"}}
            ],
            "active": True
        }
        return await db.custom_patterns.find(custom_search_query).limit(limit).to_list(length=None)
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Searching custom patterns in local fallback store: {e}")
        matches = []
        search_lower = search_term.lower()
        for pattern in list_records("custom_patterns", {"active": True}, sort_field="created_at", descending=True):
            searchable = " ".join([
                str(pattern.get("pattern", "")),
                str(pattern.get("description", "")),
                str(pattern.get("solution", "")),
                str(pattern.get("category", "")),
                str(pattern.get("example", "")),
            ]).lower()
            if search_lower in searchable or any(term and term in searchable for term in search_terms):
                matches.append(pattern)
            if len(matches) >= limit:
                break
        return matches


async def load_error_categorizations_records() -> Dict[str, List[Dict[str, Any]]]:
    try:
        permanent_patterns = await db.custom_patterns.find({
            "active": True,
            "categorization_type": "permanent"
        }).sort("created_at", -1).to_list(100)
        session_patterns = await db.session_patterns.find({
            "categorization_type": "session"
        }).sort("session_timestamp", -1).to_list(100)
        return {
            "permanent_patterns": permanent_patterns,
            "session_patterns": session_patterns,
        }
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Using local fallback for error categorizations: {e}")
        return {
            "permanent_patterns": list_records("custom_patterns", {"active": True, "categorization_type": "permanent"}, sort_field="created_at", descending=True, limit=100),
            "session_patterns": list_records("session_patterns", {"categorization_type": "session"}, sort_field="session_timestamp", descending=True, limit=100),
        }


async def save_status_check_record(record: Dict[str, Any]) -> Dict[str, Any]:
    if force_local_store:
        return insert_record("status_checks", record)

    try:
        await db.status_checks.insert_one(record)
        return record
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Saving status check to local fallback store: {e}")
        return insert_record("status_checks", record)


async def load_status_check_records(limit: int = 1000) -> List[Dict[str, Any]]:
    if force_local_store:
        return list_records("status_checks", sort_field="timestamp", descending=True, limit=limit)

    try:
        return await db.status_checks.find().sort("timestamp", -1).to_list(limit)
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Using local fallback for status checks: {e}")
        return list_records("status_checks", sort_field="timestamp", descending=True, limit=limit)


async def save_log_analysis_record(record: Dict[str, Any]) -> Dict[str, Any]:
    if force_local_store:
        return insert_record("log_analysis", record)

    try:
        await db.log_analysis.insert_one(record)
        return record
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Saving log analysis history to local fallback store: {e}")
        return insert_record("log_analysis", record)


async def load_log_analysis_records(limit: int = 50) -> List[Dict[str, Any]]:
    if force_local_store:
        return list_records("log_analysis", sort_field="timestamp", descending=True, limit=limit)

    try:
        return await db.log_analysis.find().sort("timestamp", -1).limit(limit).to_list(limit)
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Using local fallback for analysis history: {e}")
        return list_records("log_analysis", sort_field="timestamp", descending=True, limit=limit)


async def save_analysis_change_record(record: Dict[str, Any]) -> Dict[str, Any]:
    if force_local_store:
        return insert_record("analysis_changes", record)

    try:
        await db.analysis_changes.insert_one(record)
        return record
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Saving analysis changes to local fallback store: {e}")
        return insert_record("analysis_changes", record)


async def load_analysis_change_records(limit: int = 100) -> List[Dict[str, Any]]:
    if force_local_store:
        return list_records("analysis_changes", sort_field="timestamp", descending=True, limit=limit)

    try:
        return await db.analysis_changes.find().sort("timestamp", -1).to_list(limit)
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Using local fallback for analysis changes: {e}")
        return list_records("analysis_changes", sort_field="timestamp", descending=True, limit=limit)


async def find_issue_record(filters: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    if force_local_store:
        return find_record("issues", filters)

    try:
        return await db.issues.find_one(filters)
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Searching issue in local fallback store: {e}")
        return find_record("issues", filters)


async def list_issue_records(limit: Optional[int] = None) -> List[Dict[str, Any]]:
    if force_local_store:
        return list_records("issues", sort_field="data_criacao", descending=True, limit=limit)

    try:
        cursor = db.issues.find({}, {"_id": 0}).sort("data_criacao", -1)
        return await cursor.to_list(length=limit)
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Using local fallback for issues list: {e}")
        return list_records("issues", sort_field="data_criacao", descending=True, limit=limit)


async def save_issue_record(record: Dict[str, Any]) -> Dict[str, Any]:
    if force_local_store:
        return insert_record("issues", record)

    try:
        await db.issues.insert_one(record)
        return record
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Saving issue to local fallback store: {e}")
        return insert_record("issues", record)


async def update_issue_records(filters: Dict[str, Any], updates: Dict[str, Any]) -> int:
    if force_local_store:
        return update_records("issues", filters, updates)

    try:
        result = await db.issues.update_one(filters, {"$set": updates})
        return result.matched_count
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Updating issue in local fallback store: {e}")
        return update_records("issues", filters, updates)


async def delete_issue_records(filters: Dict[str, Any], limit: Optional[int] = 1) -> int:
    if force_local_store:
        return delete_records("issues", filters, limit=limit)

    try:
        result = await db.issues.delete_one(filters)
        return result.deleted_count
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Deleting issue in local fallback store: {e}")
        return delete_records("issues", filters, limit=limit)


def _normalize_auth_username(username: str) -> str:
    return username.strip()


def _normalize_auth_email(email: str) -> str:
    return email.strip().lower()


def _validate_auth_username(username: str) -> Optional[str]:
    if len(username) < 3:
        return "O usuário deve ter pelo menos 3 caracteres."

    if not re.fullmatch(r"[A-Za-z0-9._-]+", username):
        return "O usuário pode conter apenas letras, números, ponto, traço e sublinhado."

    return None


def _validate_auth_email(email: str) -> Optional[str]:
    if not re.fullmatch(r"[^\s@]+@[^\s@]+\.[^\s@]+", email):
        return "Informe um e-mail válido para concluir o acesso."

    return None


def _hash_password(password: str) -> str:
    return hashlib.sha256(password.encode("utf-8")).hexdigest()


def _validate_auth_password(password: str) -> Optional[str]:
    policy_message = "A senha deve ter no mínimo 8 caracteres, incluindo letra maiúscula, número e caractere especial."

    if len(password) < 8:
        return policy_message

    if not re.search(r"[A-ZÀ-Ý]", password):
        return policy_message

    if not re.search(r"\d", password):
        return policy_message

    if not re.search(r"[^A-Za-z0-9]", password):
        return policy_message

    return None


def _create_access_token(user: Dict[str, Any]) -> Dict[str, Any]:
    issued_at = datetime.utcnow()
    expires_at = issued_at + timedelta(minutes=JWT_EXPIRE_MINUTES)
    payload = {
        'sub': str(user.get('username', '')),
        'display_name': str(user.get('display_name', '')),
        'email': str(user.get('email', '')),
        'iat': int(issued_at.timestamp()),
        'exp': int(expires_at.timestamp()),
        'jti': str(uuid.uuid4()),
    }
    return {
        'token': jwt.encode(payload, JWT_SECRET_KEY, algorithm=JWT_ALGORITHM),
        'expires_in': JWT_EXPIRE_MINUTES * 60,
        'expires_at': expires_at.isoformat() + 'Z',
    }


def _decode_access_token(token: str) -> Optional[Dict[str, Any]]:
    try:
        payload = jwt.decode(token, JWT_SECRET_KEY, algorithms=[JWT_ALGORITHM])
    except JWTError:
        return None

    if not payload.get('sub'):
        return None

    return payload


def _sanitize_auth_user(record: Dict[str, Any]) -> Dict[str, str]:
    return {
        "username": str(record.get("username", "")),
        "display_name": str(record.get("display_name", "")),
        "email": str(record.get("email", "")),
    }


async def load_auth_users(limit: int = 5000) -> List[Dict[str, Any]]:
    if await _ensure_auth_store_ready():
        return list_records("auth_users", {"active": True}, sort_field="created_at", descending=True, limit=limit)

    if force_local_store:
        return list_records("auth_users", {"active": True}, sort_field="created_at", descending=True, limit=limit)

    try:
        return await db.auth_users.find({"active": True}).sort("created_at", -1).to_list(limit)
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Using local fallback for auth users: {e}")
        return list_records("auth_users", {"active": True}, sort_field="created_at", descending=True, limit=limit)


async def save_auth_user_record(record: Dict[str, Any]) -> Dict[str, Any]:
    if await _ensure_auth_store_ready():
        return insert_record("auth_users", record)

    if force_local_store:
        return insert_record("auth_users", record)

    try:
        await db.auth_users.insert_one(record)
        return record
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Saving auth user to local fallback store: {e}")
        return insert_record("auth_users", record)


async def find_auth_user_by_username(username: str) -> Optional[Dict[str, Any]]:
    normalized_username = _normalize_auth_username(username)

    if await _ensure_auth_store_ready():
        users = list_records("auth_users", {"active": True}, sort_field="created_at", descending=True)
        return next(
            (
                user for user in users
                if str(user.get("username_normalized", "")).lower() == normalized_username.lower()
            ),
            None,
        )

    if force_local_store:
        users = list_records("auth_users", {"active": True}, sort_field="created_at", descending=True)
        return next(
            (
                user for user in users
                if str(user.get("username_normalized", "")).lower() == normalized_username.lower()
            ),
            None,
        )

    try:
        return await db.auth_users.find_one({
            "username_normalized": normalized_username.lower(),
            "active": True,
        })
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Searching auth user by username in local fallback store: {e}")
        users = list_records("auth_users", {"active": True}, sort_field="created_at", descending=True)
        return next(
            (
                user for user in users
                if str(user.get("username_normalized", "")).lower() == normalized_username.lower()
            ),
            None,
        )


async def find_auth_user_by_email(email: str) -> Optional[Dict[str, Any]]:
    normalized_email = _normalize_auth_email(email)

    if await _ensure_auth_store_ready():
        users = list_records("auth_users", {"active": True}, sort_field="created_at", descending=True)
        return next(
            (
                user for user in users
                if str(user.get("email_normalized", "")).lower() == normalized_email
            ),
            None,
        )

    if force_local_store:
        users = list_records("auth_users", {"active": True}, sort_field="created_at", descending=True)
        return next(
            (
                user for user in users
                if str(user.get("email_normalized", "")).lower() == normalized_email
            ),
            None,
        )

    try:
        return await db.auth_users.find_one({
            "email_normalized": normalized_email,
            "active": True,
        })
    except Exception as e:
        if not _should_use_local_store(e):
            raise
        logger.warning(f"Searching auth user by email in local fallback store: {e}")
        users = list_records("auth_users", {"active": True}, sort_field="created_at", descending=True)
        return next(
            (
                user for user in users
                if str(user.get("email_normalized", "")).lower() == normalized_email
            ),
            None,
        )


@api_router.post("/auth/register", response_model=AuthResult)
async def register_auth_user(payload: AuthUserCreate):
    display_name = payload.display_name.strip()
    username = _normalize_auth_username(payload.username)
    email = _normalize_auth_email(payload.email)
    password = payload.password

    if not display_name or not username or not email or not password:
        raise fastapi.HTTPException(status_code=400, detail="Preencha todos os campos para concluir o cadastro.")

    username_error = _validate_auth_username(username)
    if username_error:
        raise fastapi.HTTPException(status_code=400, detail=username_error)

    email_error = _validate_auth_email(email)
    if email_error:
        raise fastapi.HTTPException(status_code=400, detail=email_error)

    password_error = _validate_auth_password(password)
    if password_error:
        raise fastapi.HTTPException(status_code=400, detail=password_error)

    existing_username, existing_email = await asyncio.gather(
        find_auth_user_by_username(username),
        find_auth_user_by_email(email),
    )

    if existing_username:
        raise fastapi.HTTPException(status_code=409, detail="Já existe um usuário cadastrado com esse login.")

    if existing_email:
        raise fastapi.HTTPException(status_code=409, detail="Já existe um usuário cadastrado com esse e-mail.")

    record = {
        "id": str(uuid.uuid4()),
        "display_name": display_name,
        "username": username,
        "username_normalized": username.lower(),
        "email": email,
        "email_normalized": email,
        "password_hash": _hash_password(password),
        "active": True,
        "created_at": datetime.utcnow(),
    }

    saved_user = await save_auth_user_record(record)
    return {
        "success": True,
        "message": "Usuário cadastrado com sucesso.",
        "user": _sanitize_auth_user(saved_user),
    }


@api_router.post("/auth/login", response_model=AuthResult)
async def login_auth_user(payload: AuthLoginRequest):
    username = _normalize_auth_username(payload.username)
    email = _normalize_auth_email(payload.username)
    password = payload.password

    if not username or not password:
        raise fastapi.HTTPException(status_code=400, detail="Informe usuário e senha para entrar.")

    if '@' in username:
        email_error = _validate_auth_email(email)
        if email_error:
            raise fastapi.HTTPException(status_code=400, detail=email_error)
    else:
        username_error = _validate_auth_username(username)
        if username_error:
            raise fastapi.HTTPException(status_code=400, detail=username_error)

    user = await find_auth_user_by_username(username)
    if not user and '@' in email:
        user = await find_auth_user_by_email(email)

    if not user or user.get("password_hash") != _hash_password(password):
        raise fastapi.HTTPException(status_code=401, detail="Credenciais inválidas. Solicite cadastro para acessar o analisador.")

    token_data = _create_access_token(user)

    return {
        "success": True,
        "message": "Login realizado com sucesso.",
        "user": _sanitize_auth_user(user),
        "access_token": token_data['token'],
        "token_type": "Bearer",
        "expires_in": token_data['expires_in'],
        "expires_at": token_data['expires_at'],
    }


@app.middleware('http')
async def jwt_auth_middleware(request: fastapi.Request, call_next):
    path = request.url.path

    if request.method == 'OPTIONS' or not path.startswith('/api') or path in PUBLIC_API_PATHS:
        return await call_next(request)

    authorization = request.headers.get('Authorization', '')
    if not authorization.startswith('Bearer '):
        return JSONResponse(status_code=401, content={'detail': 'Token JWT ausente ou inválido.'})

    token = authorization.replace('Bearer ', '', 1).strip()
    payload = _decode_access_token(token)
    if not payload:
        return JSONResponse(status_code=401, content={'detail': 'Token JWT expirado ou inválido.'})

    request.state.auth_user = {
        'username': payload.get('sub', ''),
        'display_name': payload.get('display_name', ''),
        'email': payload.get('email', ''),
    }
    return await call_next(request)

# Carregar padrões personalizados ao inicializar
@app.on_event("startup")
async def startup_event():
    """Carrega padrões personalizados, Datasul, LOGIX e TOTVS do banco ao inicializar o servidor."""
    try:
        use_local_runtime = await _ensure_auth_store_ready()
        startup_db = None if use_local_runtime else db

        if use_local_runtime:
            logger.warning("MongoDB indisponível no startup. Inicializando backend em modo local para evitar custo no primeiro acesso.")

        # Carregar padrões customizados
        await analyzer.load_custom_patterns_from_db(startup_db)
        logger.info("Custom patterns loaded successfully")
        
        # Inicializar carregador Datasul MongoDB
        datasul_success = await analyzer.initialize_datasul_loader(startup_db)
        if datasul_success:
            logger.info("Datasul patterns loaded from MongoDB successfully")
        else:
            logger.warning("Failed to load Datasul patterns from MongoDB")
        
        # Inicializar carregador LOGIX
        await analyzer.initialize_logix_loader(startup_db)
        logger.info("LOGIX patterns loaded successfully")
        
        # NOVO: Inicializar carregador de erros TOTVS específicos
        await analyzer.initialize_totvs_loader(startup_db)
        logger.info("TOTVS errors patterns loaded successfully")
            
    except Exception as e:
        logger.error(f"Error loading patterns: {e}")

@api_router.post("/analyze-log", response_model=LogAnalysisResult)
async def analyze_log(
    log_file: fastapi.UploadFile = fastapi.File(...),
    patterns_file: Optional[fastapi.UploadFile] = fastapi.File(None)
):
    """Analisa um arquivo de log e retorna os resultados."""
    try:
        request_started_at = monotonic()
        stage_timings = {
            "read_and_decode_ms": 0.0,
            "detect_log_type_ms": 0.0,
            "load_patterns_ms": 0.0,
            "analysis_ms": 0.0,
            "save_history_ms": 0.0,
            "total_request_ms": 0.0,
        }

        # Verificar se é um arquivo de texto
        if not log_file.content_type or not log_file.content_type.startswith('text/'):
            logger.warning(f"File type: {log_file.content_type}")
            # Continuar mesmo assim, pode ser um arquivo de log sem content-type correto
        
        # Ler conteúdo do arquivo de log
        read_started_at = monotonic()
        log_content = await log_file.read()
        log_text = decode_uploaded_text(log_content)
        stage_timings["read_and_decode_ms"] = round((monotonic() - read_started_at) * 1000, 2)

        detect_started_at = monotonic()
        detected_log_type = analyzer.detect_log_type(log_text)
        stage_timings["detect_log_type_ms"] = round((monotonic() - detect_started_at) * 1000, 2)
        logger.info(f"Pre-detected log type before analysis: {detected_log_type}")
        
        # Ler arquivo de padrões se fornecido
        patterns_content = None
        if patterns_file:
            patterns_content_bytes = await patterns_file.read()
            patterns_content = decode_uploaded_text(patterns_content_bytes)
        
        # OTIMIZAÇÃO: O analyzer já é rápido o suficiente para qualquer tamanho de log
        # graças ao pre-filter de linhas informativas
        line_count = len(log_text.split('\n'))
        logger.info(f"Using optimized standard processor for {line_count:,} lines")
        
        # Carregar padrões customizados do banco
        load_patterns_started_at = monotonic()
        custom_patterns = await load_active_custom_patterns(100)
        analyzer.load_custom_patterns([p["pattern"] for p in custom_patterns])
        
        # Carregar padrões de não-erro do banco
        non_error_patterns = await load_active_non_error_patterns(1000)
        analyzer.load_non_error_patterns([p["pattern"] for p in non_error_patterns])
        stage_timings["load_patterns_ms"] = round((monotonic() - load_patterns_started_at) * 1000, 2)
        
        # Analisar o log
        analysis_started_at = monotonic()
        result = analyzer.analyze_log_content(
            log_text, 
            patterns_content,
            enable_structured_parsing=True,
            detected_log_type=detected_log_type
        )
        stage_timings["analysis_ms"] = round((monotonic() - analysis_started_at) * 1000, 2)
        
        # Garantir que chart_data existe
        if 'chart_data' not in result or result['chart_data'] is None:
            result['chart_data'] = {
                'error_types': {'labels': [], 'values': []},
                'temporal': {'labels': [], 'values': []},
                'severity': {'labels': [], 'values': []},
                'hourly': {'labels': [], 'values': []}
            }
        
        # Salvar resultado no banco de dados para histórico
        analysis_record = {
            "id": str(uuid.uuid4()),
            "filename": log_file.filename,
            "timestamp": datetime.utcnow(),
            "log_type": detected_log_type,
            "total_results": result.get('total_results', 0),
            "statistics": result.get('statistics', {}),
            "error_counts": result.get('error_counts', {})
        }
        
        save_history_started_at = monotonic()
        await save_log_analysis_record(analysis_record)
        stage_timings["save_history_ms"] = round((monotonic() - save_history_started_at) * 1000, 2)
        stage_timings["total_request_ms"] = round((monotonic() - request_started_at) * 1000, 2)
        result['analysis_timings'] = {
            **result.get('analysis_timings', {}),
            **stage_timings,
        }
        analysis_record["analysis_timings"] = result['analysis_timings']
        logger.info(
            "Analyze-log timings for %s: detect=%sms structured=%sms patterns=%sms post=%sms total_analysis=%sms total_request=%sms",
            log_file.filename,
            result['analysis_timings'].get('detect_log_type_ms', 0.0),
            result['analysis_timings'].get('structured_parsing_ms', 0.0),
            result['analysis_timings'].get('pattern_analysis_ms', 0.0),
            result['analysis_timings'].get('post_processing_ms', 0.0),
            result['analysis_timings'].get('total_analysis_ms', 0.0),
            result['analysis_timings'].get('total_request_ms', 0.0),
        )
        
        return LogAnalysisResult(**result)
    
    except Exception as e:
        logger.error(f"Error analyzing log file: {e}")
        raise fastapi.HTTPException(status_code=500, detail=f"Erro ao analisar arquivo: {str(e)}")

@api_router.post("/download-csv")
async def download_csv(
    log_file: fastapi.UploadFile = fastapi.File(...),
    patterns_file: Optional[fastapi.UploadFile] = fastapi.File(None)
):
    """Gera e retorna um arquivo CSV com os resultados da análise."""
    try:
        # Ler e analisar o arquivo
        log_content = await log_file.read()
        log_text = decode_uploaded_text(log_content)
        detected_log_type = analyzer.detect_log_type(log_text)
        
        patterns_content = None
        if patterns_file:
            patterns_content_bytes = await patterns_file.read()
            patterns_content = decode_uploaded_text(patterns_content_bytes)
        
        # Carregar padrões customizados do banco
        custom_patterns = await load_active_custom_patterns(100)
        analyzer.load_custom_patterns([p["pattern"] for p in custom_patterns])
        
        # Carregar padrões de não-erro do banco
        non_error_patterns = await load_active_non_error_patterns(1000)
        analyzer.load_non_error_patterns([p["pattern"] for p in non_error_patterns])
        
        # Analisar
        result = analyzer.analyze_log_content(log_text, patterns_content, detected_log_type=detected_log_type)
        
        if not result.get('success', False):
            raise fastapi.HTTPException(status_code=400, detail="Erro na análise do log")
        
        # Gerar CSV
        csv_content = analyzer.generate_csv_content(result['results'])
        
        # Retornar como download
        filename = f"analise_log_{log_file.filename}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        
        # Corrigir StreamingResponse para usar io.BytesIO com bytes
        csv_bytes = csv_content.encode('utf-8-sig')  # UTF-8 com BOM para Excel
        
        return StreamingResponse(
            io.BytesIO(csv_bytes),
            media_type="text/csv; charset=utf-8",
            headers={
                "Content-Disposition": f"attachment; filename=\"{filename}\"",
                "Content-Length": str(len(csv_bytes))
            }
        )
        
    except Exception as e:
        logger.error(f"Error generating CSV: {e}")
        raise fastapi.HTTPException(status_code=500, detail=f"Erro ao gerar CSV: {str(e)}")

@api_router.post("/search-knowledge-base")
async def search_knowledge_base(request: dict):
    """Pesquisa manual na base de conhecimento de padrões de erro."""
    try:
        search_term = request.get('search_term', '').strip()
        requested_limit = request.get('max_results', 50)
        if not search_term:
            raise fastapi.HTTPException(status_code=400, detail="Termo de busca é obrigatório")

        try:
            max_results = max(1, min(int(requested_limit), 100))
        except (TypeError, ValueError):
            max_results = 50
        
        logger.info(f"Searching knowledge base for: {search_term}")
        
        # Preparar termos para busca mais flexível (palavras-chave individuais)
        search_terms = search_term.lower().split()
        # Criar regex que busca por qualquer palavra do termo
        flexible_regex = "|".join([term.replace("*", "").strip() for term in search_terms if len(term.strip()) > 2])
        
        # Pesquisar nos padrões Datasul do MongoDB
        datasul_matches = []
        try:
            # Buscar nos JSON patterns do DatasulHybridLoader
            if hasattr(analyzer, 'datasul_loader') and analyzer.datasul_loader:
                all_datasul_patterns = analyzer.datasul_loader.get_all_patterns()
                
                search_lower = search_term.lower()
                for pattern in all_datasul_patterns:
                    # Busca flexível em múltiplos campos
                    searchable = ' '.join([
                        str(pattern.get('pattern', '')),
                        str(pattern.get('description', '')),
                        str(pattern.get('solution', '')),
                        str(pattern.get('tag', '')),
                        str(pattern.get('category', '')),
                        str(pattern.get('example', ''))
                    ]).lower()
                    
                    if search_lower in searchable:
                        datasul_matches.append({
                            "type": "Padrão Datasul",
                            "code": pattern.get("tag", ""),
                            "category": pattern.get("category", ""),
                            "severity": pattern.get("severity", ""),
                            "description": pattern.get("description", ""),
                            "solution": pattern.get("solution", ""),
                            "example": pattern.get("example", ""),
                            "pattern": pattern.get("pattern", ""),
                            "source": pattern.get("source", "Datasul Knowledge Base")
                        })
                        
                        if len(datasul_matches) >= 15:
                            break
                
        except Exception as e:
            logger.error(f"Error searching Datasul patterns: {e}")
            
        # Pesquisar nos padrões PERSONALIZADOS do MongoDB
        custom_matches = []
        try:
            custom_patterns = await search_custom_patterns_records(search_term, limit=10)
            
            for pattern in custom_patterns:
                custom_matches.append({
                    "type": "Padrão Personalizado",
                    "code": pattern.get("name", "Custom"),
                    "category": pattern.get("category", "Personalizado"),
                    "severity": pattern.get("severity", "Médio"),
                    "description": pattern.get("description", "Padrão personalizado definido pelo usuário"),
                    "solution": pattern.get("solution", "Verificar contexto específico do padrão"),
                    "example": pattern.get("example", pattern.get("pattern", "")),
                    "pattern": pattern.get("pattern", ""),
                    "source": "MongoDB - Padrões Personalizados"
                })
                
        except Exception as e:
            logger.error(f"Error searching custom patterns: {e}")
        
        # Pesquisar nos padrões LOGIX
        logix_matches = []
        try:
            if hasattr(analyzer, 'logix_loader') and analyzer.logix_loader:
                logix_results = analyzer.logix_loader.search_patterns(search_term, limit=15)
                
                for pattern in logix_results:
                    logix_matches.append({
                        "type": "Padrão LOGIX",
                        "code": pattern.get("tag", ""),
                        "category": pattern.get("category", ""),
                        "severity": pattern.get("severity", ""),
                        "description": pattern.get("description", ""),
                        "solution": pattern.get("solution", ""),
                        "example": pattern.get("example", ""),
                        "pattern": pattern.get("pattern", ""),
                        "source": pattern.get("source", "LOGIX Knowledge Base")
                    })
        except Exception as e:
            logger.error(f"Error searching LOGIX patterns: {e}")

        # Pesquisar nos padrões TOTVS específicos
        totvs_matches = []
        try:
            if hasattr(analyzer, 'totvs_loader') and analyzer.totvs_loader:
                totvs_results = analyzer.totvs_loader.search_patterns(search_term, limit=15)

                for pattern in totvs_results:
                    totvs_matches.append({
                        "type": "Padrão TOTVS",
                        "code": pattern.get("code", pattern.get("tag", "")),
                        "category": pattern.get("category", ""),
                        "severity": pattern.get("severity", ""),
                        "description": pattern.get("description", ""),
                        "solution": pattern.get("solution", ""),
                        "example": pattern.get("example", ""),
                        "pattern": pattern.get("pattern", ""),
                        "source": pattern.get("reference", pattern.get("source", "TOTVS Knowledge Base"))
                    })
        except Exception as e:
            logger.error(f"Error searching TOTVS patterns: {e}")
        
        # Pesquisar nos padrões hardcoded Progress/OpenEdge/PASOE
        hardcoded_matches = []
        
        # Padrões Progress/OpenEdge conhecidos expandidos
        progress_errors = [
            {
                "code": "1793",
                "category": "OpenEdge/DB", 
                "severity": "Alto",
                "description": "Falha na conexão com banco de dados Progress",
                "solution": "Verificar conectividade de rede, parâmetros de conexão (-H, -S, -db), e status do database broker",
                "example": "Erro 1793 - Database connection failed"
            },
            {
                "code": "204",
                "category": "Datasul/NFe",
                "severity": "Alto", 
                "description": "Rejeição 204 - Duplicidade de NF-e",
                "solution": "Verificar se NF-e já foi transmitida. Cancelar duplicata ou usar número diferente",
                "example": "Rejeição 204 - Duplicidade de NF-e"
            },
            {
                "code": "NFe already exists",
                "category": "Datasul/NFe",
                "severity": "Alto", 
                "description": "Parâmetro NFe já existe para empresa - erro de duplicidade",
                "solution": "Verificar se parâmetro NFe já foi configurado para essa empresa. Remover duplicata ou usar empresa diferente",
                "example": "** Parametro NFe already exists with Empresa 000001"
            },
            {
                "code": "Parametro NFe",
                "category": "Datasul/NFe",
                "severity": "Médio", 
                "description": "Erro relacionado a configuração de parâmetros NFe no Datasul",
                "solution": "Verificar configuração dos parâmetros NFe, validar empresa e certificado digital",
                "example": "Parametro NFe already exists with Empresa"
            },
            {
                "code": "PASOE",
                "category": "AppServer/PASOE",
                "severity": "Crítico",
                "description": "Falha no Progress Application Server OpenEdge",
                "solution": "Verificar logs do PASOE, reiniciar serviços, verificar configuração de instâncias",
                "example": "PASOE instance failed to start"
            },
            {
                "code": "Procedure:",
                "category": "OpenEdge/ABL",
                "severity": "Médio",
                "description": "Erro em execução de procedure Progress 4GL",
                "solution": "Verificar sintaxe, parâmetros e dependências da procedure. Consultar stack trace",
                "example": "Procedure: /usr/dlc/test.p failed"
            },
            {
                "code": "LOG:MANAGER",
                "category": "Datasul/Sistema",
                "severity": "Alto",
                "description": "Erro no gerenciador de logs Datasul",
                "solution": "Verificar permissões de escrita, espaço em disco e configuração do log manager",
                "example": "LOG:MANAGER - Write permission denied"
            },
            {
                "code": "AppServer",
                "category": "AppServer/PASOE",
                "severity": "Alto",
                "description": "Falha no Application Server Progress",
                "solution": "Reiniciar serviços AppServer, verificar configurações de conexão e recursos do sistema",
                "example": "AppServer connection lost"
            },
            {
                "code": "already exists",
                "category": "Datasul/Sistema",
                "severity": "Médio",
                "description": "Erro de duplicidade - registro já existe no sistema",
                "solution": "Verificar se registro já foi criado. Remover duplicata ou usar chave única diferente",
                "example": "Record already exists in database"
            }
        ]
        
        # Filtrar padrões hardcoded com busca MUITO flexível
        search_lower = search_term.lower()
        search_words = [word.strip() for word in search_lower.split() if len(word.strip()) > 2]
        
        for error in progress_errors:
            # Busca exata
            exact_match = (search_lower in error["code"].lower() or 
                          search_lower in error["description"].lower() or
                          search_lower in error["solution"].lower() or
                          search_lower in error["category"].lower() or
                          search_lower in error["example"].lower())
            
            # Busca por palavras individuais
            word_match = False
            if search_words:
                for word in search_words:
                    if (word in error["code"].lower() or 
                        word in error["description"].lower() or
                        word in error["solution"].lower() or
                        word in error["example"].lower()):
                        word_match = True
                        break
            
            if exact_match or word_match:
                hardcoded_matches.append({
                    "type": "Padrão Sistema",
                    "code": error["code"],
                    "category": error["category"],
                    "severity": error["severity"],
                    "description": error["description"],
                    "solution": error["solution"],
                    "example": error["example"],
                    "pattern": "",
                    "source": "Sistema - Padrões Progress/PASOE"
                })
        
        # Combinar todos os tipos de padrões
        all_matches = datasul_matches + logix_matches + totvs_matches + custom_matches + hardcoded_matches
        
        # Ordenar por relevância (matches exatos primeiro)
        def match_score(match):
            score = 0
            if search_lower in match["code"].lower():
                score += 10
            if search_lower in match["description"].lower():
                score += 5
            if search_lower in match["solution"].lower():
                score += 3
            return score
        
        all_matches.sort(key=match_score, reverse=True)
        
        # Retornar até o limite solicitado, preservando o total encontrado.
        final_matches = all_matches[:max_results]
        returned_count = len(final_matches)
        total_found = len(all_matches)
        
        return {
            "success": True,
            "search_term": search_term,
            "total_found": total_found,
            "returned_count": returned_count,
            "truncated": total_found > returned_count,
            "max_results": max_results,
            "matches": final_matches,
            "sources": {
                "datasul_patterns": len(datasul_matches),
                "logix_patterns": len(logix_matches),
                "totvs_patterns": len(totvs_matches),
                "custom_patterns": len(custom_matches),
                "system_patterns": len(hardcoded_matches)
            }
        }
        
    except Exception as e:
        logger.error(f"Error searching knowledge base: {e}")
        raise fastapi.HTTPException(status_code=500, detail=f"Erro na pesquisa: {str(e)}")

@api_router.post("/analyze-log-categories")
async def analyze_log_categories(
    log_file: fastapi.UploadFile = fastapi.File(...),
    selected_log_type: str = fastapi.Form("auto")
):
    """Analisa as categorias/famílias presentes no log suportado pelo limpador."""
    try:
        from log_cleaner import (
            LogCleaner,
            build_grouped_category_matches,
            get_categories_for_log_type,
            get_cleaner_log_type_label,
            resolve_cleaner_log_type,
        )
        
        # Ler o arquivo
        content = await log_file.read()
        try:
            log_text = content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                log_text = content.decode('latin-1')
            except UnicodeDecodeError:
                log_text = content.decode('utf-8', errors='ignore')
        
        # Analisar categorias
        cleaner = LogCleaner()
        analysis = cleaner.analyze_log(log_text)
        detected_log_type = analysis.get("format_info", {}).get("format", "generic")
        normalized_selected_log_type, effective_log_type = resolve_cleaner_log_type(selected_log_type, detected_log_type)
        allowed_categories = get_categories_for_log_type(effective_log_type)
        category_info = cleaner.get_category_info(effective_log_type)
        found_groups = build_grouped_category_matches(analysis, allowed_categories=allowed_categories)
        
        # Filtrar apenas categorias que foram encontradas
        found_categories = {}
        for category, count in analysis["totals"].items():
            if category not in allowed_categories:
                continue
            if category not in ["_considered", "_unmatched"] and count > 0:
                found_categories[category] = {
                    "count": count,
                    "display_name": category_info["display_names"].get(category, category),
                    "samples": analysis["samples"].get(category, [])[:3]  # Máximo 3 amostras
                }
        
        return {
            "success": True,
            "filename": log_file.filename,
            "file_size": len(content),
            "total_lines": len(log_text.splitlines()),
            "analysis": analysis["totals"],
            "selected_log_type": normalized_selected_log_type,
            "selected_log_type_label": get_cleaner_log_type_label(normalized_selected_log_type),
            "detected_log_type": detected_log_type,
            "detected_log_type_label": get_cleaner_log_type_label(detected_log_type),
            "effective_log_type": effective_log_type,
            "effective_log_type_label": get_cleaner_log_type_label(effective_log_type),
            "found_categories": found_categories,
            "found_groups": found_groups,
            "category_info": category_info,
            "samples": analysis["samples"]
        }
        
    except Exception as e:
        logger.error(f"Error analyzing log categories: {e}")
        raise fastapi.HTTPException(status_code=500, detail=f"Erro ao analisar categorias: {str(e)}")

@api_router.post("/clean-log")
async def clean_log(
    log_file: fastapi.UploadFile = fastapi.File(...),
    categories_to_remove: str = fastapi.Form("")
):
    """Limpa o log removendo as categorias selecionadas."""
    try:
        from log_cleaner import LogCleaner
        
        # Ler o arquivo
        content = await log_file.read()
        try:
            log_text = content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                log_text = content.decode('latin-1')
            except UnicodeDecodeError:
                log_text = content.decode('utf-8', errors='ignore')
        
        # Parsear categorias a serem removidas
        try:
            categories_list = categories_to_remove.split(",") if categories_to_remove else []
            categories_list = [cat.strip() for cat in categories_list if cat.strip()]
        except:
            categories_list = []
        
        if not categories_list:
            raise fastapi.HTTPException(status_code=400, detail="Nenhuma categoria selecionada para remoção")
        
        # Limpar o log
        cleaner = LogCleaner()
        result = cleaner.clean_log(log_text, categories_list)
        
        if not result["success"]:
            raise fastapi.HTTPException(status_code=400, detail="Erro ao limpar o log")
        
        # Preparar arquivo para download
        cleaned_content = result["cleaned_content"]
        filename_base = log_file.filename.rsplit('.', 1)[0] if '.' in log_file.filename else log_file.filename
        cleaned_filename = f"{filename_base}_limpo.log"
        
        # Retornar como download
        csv_bytes = cleaned_content.encode('utf-8-sig')  # UTF-8 com BOM
        
        return StreamingResponse(
            io.BytesIO(csv_bytes),
            media_type="text/plain; charset=utf-8",
            headers={
                "Content-Disposition": f"attachment; filename=\"{cleaned_filename}\"",
                "Content-Length": str(len(csv_bytes)),
                "X-Cleaning-Stats": json.dumps(result["statistics"])
            }
        )
        
    except Exception as e:
        logger.error(f"Error cleaning log: {e}")
        raise fastapi.HTTPException(status_code=500, detail=f"Erro ao limpar log: {str(e)}")


@api_router.post("/split-log")
async def split_log_file(
    log_file: fastapi.UploadFile = fastapi.File(...),
    lines_per_chunk: int = fastapi.Form(...)
):
    """Divide um arquivo grande de log/texto em vários arquivos menores e retorna um zip."""
    try:
        allowed_chunk_sizes = {50000, 100000, 150000, 200000, 250000}
        if lines_per_chunk not in allowed_chunk_sizes:
            raise fastapi.HTTPException(
                status_code=400,
                detail="Escolha uma quantidade válida de linhas: 50, 100, 150, 200 ou 250 mil."
            )

        content = await log_file.read()
        if not content:
            raise fastapi.HTTPException(status_code=400, detail="O arquivo enviado está vazio.")

        log_text = decode_uploaded_text(content)
        all_lines = log_text.splitlines(keepends=True)
        if not all_lines:
            raise fastapi.HTTPException(status_code=400, detail="Não foi possível identificar linhas válidas no arquivo.")

        filename = log_file.filename or "log.txt"
        filename_base = filename.rsplit('.', 1)[0] if '.' in filename else filename
        extension = filename.rsplit('.', 1)[1] if '.' in filename else 'txt'
        total_parts = (len(all_lines) + lines_per_chunk - 1) // lines_per_chunk

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, mode='w', compression=zipfile.ZIP_DEFLATED) as zip_file:
            for index, start in enumerate(range(0, len(all_lines), lines_per_chunk), start=1):
                part_lines = all_lines[start:start + lines_per_chunk]
                part_filename = f"{filename_base}_parte_{index:03d}.{extension}"
                zip_file.writestr(part_filename, ''.join(part_lines).encode('utf-8-sig'))

        zip_bytes = zip_buffer.getvalue()
        archive_filename = f"{filename_base}_dividido_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        split_stats = json.dumps({
            "original_filename": filename,
            "total_lines": len(all_lines),
            "lines_per_chunk": lines_per_chunk,
            "generated_files": total_parts,
        })

        return StreamingResponse(
            io.BytesIO(zip_bytes),
            media_type="application/zip",
            headers={
                "Content-Disposition": f"attachment; filename=\"{archive_filename}\"",
                "Content-Length": str(len(zip_bytes)),
                "X-Split-Stats": split_stats,
            }
        )

    except fastapi.HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error splitting log: {e}")
        raise fastapi.HTTPException(status_code=500, detail=f"Erro ao dividir log: {str(e)}")

@api_router.post("/analyze-profiler")
async def analyze_profiler_file(log_file: fastapi.UploadFile = fastapi.File(...)):
    """Analisa arquivo .out do Progress Profiler"""
    try:
        from profiler_analyzer import ProfilerAnalyzer
        
        content = await log_file.read()

        try:
            file_content = content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                file_content = content.decode("latin-1")
            except UnicodeDecodeError:
                file_content = content.decode("utf-8", errors="ignore")

        analyzer = ProfilerAnalyzer()
        result = analyzer.analyze_file_content(file_content)

        if not result["success"]:
            raise fastapi.HTTPException(
                status_code=400,
                detail=result["error"]
            )

        # 🔥 NOVO RETORNO MODERNO
        return {
            "success": True,
            "filename": log_file.filename,
            "file_size": len(content),
            "session": result.get("session"),
            "summary": result.get("summary"),
            "top_bottlenecks": result.get("top_bottlenecks"),
            "n_plus_one_suspects": result.get("n_plus_one_suspects"),
            "call_tree": result.get("call_tree"),
            "raw_data": result.get("raw_data"),
            "analysis": result.get("analysis")
        }

    except Exception as e:
        logger.error(f"Error analyzing profiler file: {e}")
        raise fastapi.HTTPException(
            status_code=500,
            detail=f"Erro ao analisar profiler: {str(e)}"
        )
@api_router.get("/analysis-history", response_model=List[dict])
async def get_analysis_history():
    """Retorna o histórico de análises realizadas."""
    try:
        analyses = await load_log_analysis_records(50)
        
        # Converter ObjectId para string e remover _id
        for analysis in analyses:
            analysis.pop('_id', None)
        
        return analyses
    except Exception as e:
        logger.error(f"Error fetching analysis history: {e}")
        raise fastapi.HTTPException(status_code=500, detail="Erro ao buscar histórico")

@api_router.post("/add-pattern")
async def add_custom_pattern(pattern_data: AddPatternRequest):
    """Adiciona um novo padrão de erro personalizado ao sistema com validação avançada."""
    try:
        # Validar dados básicos
        if not pattern_data.pattern.strip():
            raise fastapi.HTTPException(status_code=400, detail="Padrão não pode estar vazio")
        
        if not pattern_data.description.strip():
            raise fastapi.HTTPException(status_code=400, detail="Descrição é obrigatória")
            
        if not pattern_data.solution.strip():
            raise fastapi.HTTPException(status_code=400, detail="Possível solução é obrigatória")
        
        # NOVA VALIDAÇÃO AVANÇADA
        from pattern_validator import validate_pattern_for_api
        
        # Preparar dados para validação
        validation_data = {
            "pattern": pattern_data.pattern.strip(),
            "partial_pattern": pattern_data.partial_pattern.strip() if pattern_data.partial_pattern else "",
            "example": pattern_data.example or "",
            "description": pattern_data.description.strip()
        }
        
        # Preparar logs de teste (usando exemplo + padrões genéricos)
        test_logs = []
        if pattern_data.example:
            test_logs.append(pattern_data.example.strip())
        
        # Adicionar alguns cenários de teste genéricos
        test_logs.extend([
            pattern_data.pattern.strip(),  # O próprio padrão
            f"ERROR: {pattern_data.pattern.strip()}",  # Com prefixo
            f"[WARN] {pattern_data.pattern.strip()} occurred",  # Com contexto
            f"CRITICAL {pattern_data.pattern.strip()} detected in system"  # Mais contexto
        ])
        
        # Executar validação
        validation_result = validate_pattern_for_api(validation_data, test_logs)
        
        # Se não passou na validação básica, rejeitar
        if not validation_result["overall_valid"]:
            errors = validation_result["pattern_validation"].get("errors", [])
            raise fastapi.HTTPException(
                status_code=400, 
                detail=f"Padrão inválido: {'; '.join(errors)}"
            )
        
        # Se funcionalidade não foi testada com sucesso, avisar mas permitir
        functionality_warnings = []
        if not validation_result["functionality_test"].get("pattern_works", True):
            functionality_warnings.append("Padrão pode ter baixa taxa de detecção")
        
        # Log da validação
        match_rate = validation_result["functionality_test"].get("match_rate", 0)
        logger.info(f"Pattern validation - Match rate: {match_rate}%, Complexity: {validation_result['pattern_validation'].get('complexity', 'unknown')}")
        
        # Criar registro do novo padrão com informações de validação
        new_pattern = {
            "id": str(uuid.uuid4()),
            "pattern": pattern_data.pattern.strip(),
            "partial_pattern": pattern_data.partial_pattern.strip() if pattern_data.partial_pattern else "",
            "description": pattern_data.description.strip(),
            "category": pattern_data.category,
            "severity": pattern_data.severity,
            "example": pattern_data.example or "",
            "solution": pattern_data.solution.strip(),
            "created_at": datetime.utcnow(),
            "active": True,
            "user_created": True,
            "validation_info": {
                "match_rate": validation_result["functionality_test"].get("match_rate", 0),
                "complexity": validation_result["pattern_validation"].get("complexity", "unknown"),
                "pattern_type": validation_result["pattern_validation"].get("pattern_type", "unknown"),
                "validated_at": datetime.utcnow().isoformat(),
                "test_logs_count": len(test_logs),
                "warnings": validation_result["recommendations"] + functionality_warnings
            }
        }
        
        # Salvar no banco de dados
        await save_custom_pattern_record(new_pattern)
        
        # Adicionar ao analisador em tempo real
        analyzer.add_custom_pattern(pattern_data.pattern)
        
        logger.info(f"New custom pattern added: {pattern_data.pattern}")
        
        # Remover _id do MongoDB para retorno
        new_pattern.pop('_id', None)
        
        return {
            "success": True,
            "message": "Padrão adicionado com sucesso",
            "pattern": new_pattern,
            "validation_result": {
                "match_rate": validation_result["functionality_test"].get("match_rate", 0),
                "pattern_works": validation_result["functionality_test"].get("pattern_works", False),
                "complexity": validation_result["pattern_validation"].get("complexity", "unknown"),
                "recommendations": validation_result["recommendations"],
                "warnings": functionality_warnings,
                "suggestions": validation_result["final_suggestions"]
            }
        }
        
    except Exception as e:
        logger.error(f"Error adding custom pattern: {e}")
        raise fastapi.HTTPException(status_code=500, detail=f"Erro ao adicionar padrão: {str(e)}")

@api_router.post("/analyze-info")
async def get_log_analysis_info(log_file: fastapi.UploadFile = fastapi.File(...)):
    """Fornece informações sobre processamento do log antes da análise completa."""
    try:
        # Ler apenas uma amostra para análise rápida
        log_content = await log_file.read()
        
        # Decode content
        try:
            log_text = log_content.decode('utf-8')
        except UnicodeDecodeError:
            log_text = log_content.decode('latin1', errors='ignore')
        
        # Estatísticas básicas
        lines = log_text.split('\n')
        line_count = len(lines)
        file_size_bytes = len(log_content)
        file_size_mb = file_size_bytes / (1024 * 1024)
        
        # Amostra das primeiras linhas para preview
        sample_lines = [line.strip() for line in lines[:10] if line.strip()]
        
        # Determinar tipo de processamento
        processing_type = "standard"
        if line_count > 50000:
            processing_type = "massive"
        elif line_count > 20000:
            processing_type = "very_large"
        elif line_count > 10000:
            processing_type = "large"
        
        # Estimar tempo
        estimated_seconds = max(5, line_count / 1500)  # ~1500 linhas por segundo
        
        return {
            "success": True,
            "file_info": {
                "filename": log_file.filename,
                "size_bytes": file_size_bytes,
                "size_mb": round(file_size_mb, 2),
                "line_count": line_count,
                "processing_type": processing_type
            },
            "processing_estimate": {
                "estimated_time_seconds": round(estimated_seconds, 1),
                "estimated_time_human": f"{estimated_seconds/60:.1f} min" if estimated_seconds > 60 else f"{estimated_seconds:.0f}s",
                "will_use_optimization": line_count > 10000 or file_size_mb > 5,
                "chunk_processing": line_count > 10000
            },
            "sample_preview": {
                "first_lines": sample_lines,
                "total_preview_lines": len(sample_lines)
            },
            "recommendations": [
                f"Log detectado como '{processing_type}' ({line_count:,} linhas)",
                "Processamento otimizado será usado" if line_count > 10000 else "Processamento padrão será usado",
                "Resultados podem ser limitados para performance" if line_count > 50000 else "Análise completa disponível",
                f"Tempo estimado: {estimated_seconds/60:.1f} minutos" if estimated_seconds > 60 else f"Tempo estimado: {estimated_seconds:.0f} segundos"
            ]
        }
        
    except Exception as e:
        logger.error(f"Error getting log info: {e}")
        raise fastapi.HTTPException(status_code=500, detail=f"Erro ao analisar informações do log: {str(e)}")

@api_router.post("/test-pattern")
async def test_pattern_before_saving(request: dict):
    """Testa um padrão antes de salvá-lo para verificar se funciona corretamente."""
    try:
        from pattern_validator import validate_pattern_for_api
        
        # Extrair dados da requisição
        pattern = request.get("pattern", "").strip()
        example_log = request.get("example", "").strip()
        test_logs_provided = request.get("test_logs", [])
        
        if not pattern:
            raise fastapi.HTTPException(status_code=400, detail="Padrão é obrigatório para teste")
        
        # Preparar dados para validação
        validation_data = {
            "pattern": pattern,
            "partial_pattern": request.get("partial_pattern", ""),
            "example": example_log,
            "description": request.get("description", "")
        }
        
        # Preparar logs de teste
        test_logs = test_logs_provided.copy() if test_logs_provided else []
        
        if example_log and example_log not in test_logs:
            test_logs.append(example_log)
        
        # Se não há logs de teste, criar alguns genéricos
        if not test_logs:
            test_logs = [
                pattern,
                f"ERROR: {pattern}",
                f"[WARN] {pattern} detected",
                f"CRITICAL {pattern} in system",
                f"INFO: {pattern} occurred at 10:30"
            ]
        
        # Executar validação
        validation_result = validate_pattern_for_api(validation_data, test_logs)
        
        # Retornar resultado detalhado
        return {
            "success": True,
            "validation_result": validation_result,
            "test_summary": {
                "pattern_valid": validation_result["overall_valid"],
                "match_rate": validation_result["functionality_test"].get("match_rate", 0),
                "matches_found": validation_result["functionality_test"].get("matches_found", 0),
                "total_tests": validation_result["functionality_test"].get("total_tests", 0),
                "pattern_works": validation_result["functionality_test"].get("pattern_works", False)
            },
            "recommendations": validation_result["recommendations"],
            "suggestions": validation_result["final_suggestions"]
        }
        
    except Exception as e:
        logger.error(f"Error testing pattern: {e}")
        raise fastapi.HTTPException(status_code=500, detail=f"Erro ao testar padrão: {str(e)}")

@api_router.get("/custom-patterns")
async def get_custom_patterns():
    """Retorna todos os padrões personalizados ativos."""
    try:
        patterns = await load_active_custom_patterns(1000)
        
        # Converter ObjectId para string e remover _id
        for pattern in patterns:
            pattern.pop('_id', None)
        
        return {
            "success": True,
            "patterns": patterns
        }
        
    except Exception as e:
        logger.error(f"Error fetching custom patterns: {e}")
        raise fastapi.HTTPException(status_code=500, detail="Erro ao buscar padrões personalizados")

@api_router.delete("/custom-patterns/{pattern_id}")
async def delete_custom_pattern(pattern_id: str):
    """Remove um padrão personalizado."""
    try:
        matched_count = await deactivate_custom_pattern_record(pattern_id)
        
        if matched_count == 0:
            raise fastapi.HTTPException(status_code=404, detail="Padrão não encontrado")
        
        return {
            "success": True,
            "message": "Padrão removido com sucesso"
        }
        
    except Exception as e:
        logger.error(f"Error deleting custom pattern: {e}")
        raise fastapi.HTTPException(status_code=500, detail="Erro ao remover padrão")

@api_router.post("/search-log", response_model=SearchResult)
async def search_log(
    log_file: fastapi.UploadFile = fastapi.File(...),
    search_pattern: str = fastapi.Form(...),
    case_sensitive: bool = fastapi.Form(False),
    search_type: str = fastapi.Form("procedure")
):
    """Realiza busca específica em arquivo de log."""
    try:
        # Ler conteúdo do arquivo de log
        log_content = await log_file.read()
        try:
            log_text = log_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                log_text = log_content.decode('latin-1')
            except UnicodeDecodeError:
                log_text = log_content.decode('utf-8', errors='ignore')
        
        # Dividir em linhas
        lines = log_text.split('\n')
        matches = []
        
        # Preparar padrão de busca
        flags = 0 if case_sensitive else re.IGNORECASE
        
        try:
            # Se for busca personalizada, tratar como regex
            if search_type == "custom":
                pattern = re.compile(search_pattern, flags)
            else:
                # Para outros tipos, busca literal
                pattern = re.compile(re.escape(search_pattern), flags)
        except re.error:
            # Se regex inválido, fazer busca literal
            pattern = re.compile(re.escape(search_pattern), flags)
        
        # Buscar em cada linha
        for line_num, line in enumerate(lines, start=1):
            if pattern.search(line):
                # Extrair timestamp se possível
                timestamp = extract_timestamp_from_line(line)
                
                # Destacar o match
                highlighted_content = pattern.sub(
                    lambda m: f"**{m.group()}**", 
                    line
                ) if search_type != "custom" else line
                
                match_info = {
                    "line_number": line_num,
                    "content": line,
                    "highlighted_content": highlighted_content,
                    "timestamp": timestamp,
                    "match_position": pattern.search(line).start() if pattern.search(line) else 0
                }
                matches.append(match_info)
        
        # Informações da busca
        search_info = {
            "pattern": search_pattern,
            "search_type": search_type,
            "case_sensitive": case_sensitive,
            "total_lines_searched": len(lines),
            "filename": log_file.filename
        }
        
        logger.info(f"Search completed: {len(matches)} matches found for pattern '{search_pattern}'")
        
        return SearchResult(
            success=True,
            total_matches=len(matches),
            matches=matches[:1000],  # Limitar para performance
            search_info=search_info
        )
        
    except Exception as e:
        logger.error(f"Error searching log: {e}")
        return SearchResult(
            success=False,
            total_matches=0,
            matches=[],
            search_info={},
            error=str(e)
        )

@api_router.post("/categorize-error")
async def categorize_error(categorization: ErrorCategorization):
    """Permite ao usuário decidir se um erro deve ser permanente ou apenas da sessão."""
    try:
        if categorization.category_type == "permanent":
            # Adicionar como padrão permanente no banco
            new_pattern = {
                "id": str(uuid.uuid4()),
                "pattern": categorization.pattern,
                "description": categorization.description or "Padrão adicionado via categorização",
                "category": "User-Categorized",
                "severity": "Médio",
                "created_at": datetime.utcnow(),
                "active": True,
                "user_created": True,
                "categorization_type": "permanent"
            }
            
            await save_custom_pattern_record(new_pattern)
            analyzer.add_custom_pattern(categorization.pattern)
            
            logger.info(f"Pattern added as permanent: {categorization.pattern}")
            
            return {
                "success": True,
                "message": "Padrão adicionado permanentemente ao sistema",
                "type": "permanent",
                "pattern_id": new_pattern["id"]
            }
            
        elif categorization.category_type == "session":
            # Salvar apenas como registro de sessão (não adiciona aos padrões ativos)
            session_record = {
                "id": str(uuid.uuid4()),
                "pattern": categorization.pattern,
                "description": categorization.description or "Padrão considerado apenas para esta sessão",
                "session_timestamp": datetime.utcnow(),
                "categorization_type": "session",
                "active": False  # Não ativo para análises futuras
            }
            
            await save_session_pattern_record(session_record)
            
            logger.info(f"Pattern marked as session-only: {categorization.pattern}")
            
            return {
                "success": True,
                "message": "Padrão considerado apenas para esta análise",
                "type": "session",
                "pattern_id": session_record["id"]
            }
        else:
            raise fastapi.HTTPException(status_code=400, detail="Tipo de categoria inválido")
            
    except Exception as e:
        logger.error(f"Error categorizing pattern: {e}")
        raise fastapi.HTTPException(status_code=500, detail=f"Erro ao categorizar padrão: {str(e)}")

@api_router.get("/error-categorizations")
async def get_error_categorizations():
    """Retorna histórico de categorizações de erro."""
    try:
        categorizations = await load_error_categorizations_records()
        permanent_patterns = categorizations["permanent_patterns"]
        session_patterns = categorizations["session_patterns"]
        
        # Limpar _id do MongoDB
        for pattern in permanent_patterns + session_patterns:
            pattern.pop('_id', None)
        
        return {
            "success": True,
            "permanent_patterns": permanent_patterns,
            "session_patterns": session_patterns
        }
        
    except Exception as e:
        logger.error(f"Error fetching categorizations: {e}")
        raise fastapi.HTTPException(status_code=500, detail="Erro ao buscar categorizações")

@api_router.post("/mark-as-non-error")
async def mark_as_non_error(non_error_data: NonErrorPattern):
    """Marca um padrão como não sendo um erro para análises futuras."""
    try:
        # Salvar padrão de não-erro no banco
        non_error_record = {
            "id": str(uuid.uuid4()),
            "pattern": (non_error_data.partial_pattern or non_error_data.pattern).strip(),
            "full_message": non_error_data.full_message,
            "partial_pattern": (non_error_data.partial_pattern or non_error_data.pattern).strip(),
            "reason": non_error_data.reason or "Marcado pelo usuário como não sendo erro",
            "source_line": non_error_data.source_line,
            "created_at": datetime.utcnow(),
            "active": True
        }
        
        await save_non_error_pattern_record(non_error_record)
        
        logger.info(f"Pattern marked as non-error: {non_error_data.pattern}")
        
        return {
            "success": True,
            "message": "Padrão marcado como não-erro com sucesso",
            "pattern_id": non_error_record["id"]
        }
        
    except Exception as e:
        logger.error(f"Error marking pattern as non-error: {e}")
        raise fastapi.HTTPException(status_code=500, detail=f"Erro ao marcar como não-erro: {str(e)}")

@api_router.get("/non-error-patterns")
async def get_non_error_patterns():
    """Retorna todos os padrões marcados como não-erro."""
    try:
        patterns = await load_active_non_error_patterns(1000)
        
        for pattern in patterns:
            pattern.pop('_id', None)
        
        return {
            "success": True,
            "patterns": patterns
        }
        
    except Exception as e:
        logger.error(f"Error fetching non-error patterns: {e}")
        raise fastapi.HTTPException(status_code=500, detail="Erro ao buscar padrões de não-erro")

@api_router.post("/save-analysis-changes")
async def save_analysis_changes(changes: dict):
    """Salva alterações feitas na análise de erros."""
    try:
        # Criar registro de alterações
        change_record = {
            "id": str(uuid.uuid4()),
            "changes": changes,
            "timestamp": datetime.utcnow(),
            "user_id": "system",  # Pode ser expandido para suporte a usuários
            "change_type": changes.get("type", "unknown")
        }
        
        await save_analysis_change_record(change_record)
        
        logger.info(f"Analysis changes saved: {changes.get('type', 'unknown')}")
        
        return {
            "success": True,
            "message": "Alterações salvas com sucesso",
            "change_id": change_record["id"]
        }
        
    except Exception as e:
        logger.error(f"Error saving analysis changes: {e}")
        raise fastapi.HTTPException(status_code=500, detail=f"Erro ao salvar alterações: {str(e)}")

@api_router.get("/analysis-changes")
async def get_analysis_changes():
    """Retorna histórico de alterações da análise."""
    try:
        changes = await load_analysis_change_records(100)
        
        for change in changes:
            change.pop('_id', None)
        
        return {
            "success": True,
            "changes": changes
        }
        
    except Exception as e:
        logger.error(f"Error fetching analysis changes: {e}")
        raise fastapi.HTTPException(status_code=500, detail="Erro ao buscar histórico de alterações")

@api_router.get("/datasul-patterns")
async def get_datasul_patterns():
    """Retorna todos os padrões Datasul do MongoDB."""
    try:
        if not analyzer.datasul_loader:
            return {"success": False, "error": "Datasul loader not initialized"}
            
        patterns = analyzer.datasul_loader.get_all_patterns_with_solutions()
        patterns_list = []
        
        for pattern in patterns:
            patterns_list.append({
                "id": pattern.get("id", ""),
                "pattern": pattern.get("pattern", ""),
                "description": pattern.get("description", ""),
                "category": pattern.get("category", ""),
                "severity": pattern.get("severity", ""),
                "solution": pattern.get("solution", ""),
                "tag": pattern.get("tag", ""),
                "priority": pattern.get("priority", 3),
                "active": pattern.get("active", True),
                "usage_count": pattern.get("usage_count", 0),
                "created_at": pattern.get("created_at", ""),
                "last_detected": pattern.get("last_detected")
            })
        
        return {"success": True, "patterns": patterns_list, "total": len(patterns_list)}
        
    except Exception as e:
        logger.error(f"Error fetching Datasul patterns: {e}")
        return {"success": False, "error": str(e)}

@api_router.get("/datasul-statistics")
async def get_datasul_statistics():
    """Retorna estatísticas dos padrões Datasul."""
    try:
        if not analyzer.datasul_loader:
            return {"success": False, "error": "Datasul loader not initialized"}
            
        stats = analyzer.datasul_loader.get_statistics()
        return {"success": True, "statistics": stats}
        
    except Exception as e:
        logger.error(f"Error fetching Datasul statistics: {e}")
        return {"success": False, "error": str(e)}

@api_router.post("/refresh-datasul-patterns")
async def refresh_datasul_patterns():
    """Recarrega padrões Datasul do MongoDB (limpa cache)."""
    try:
        if not analyzer.datasul_loader:
            return {"success": False, "error": "Datasul loader not initialized"}
            
        success = await analyzer.datasul_loader.load_patterns_from_db()
        if success:
            return {"success": True, "message": "Datasul patterns refreshed successfully"}
        else:
            return {"success": False, "error": "Failed to refresh patterns from database"}
            
    except Exception as e:
        logger.error(f"Error refreshing Datasul patterns: {e}")
        return {"success": False, "error": str(e)}


@api_router.get("/version-compare/status")
async def get_version_compare_status():
    """Retorna metadados do índice do comparador de versões."""
    return {
        "success": True,
        **version_compare_service.get_index_metadata(),
    }


@api_router.post("/version-compare/reload")
async def reload_version_compare_index():
    """Recarrega o índice do comparador de versões em memória."""
    try:
        version_compare_service.reload_index()
        return {
            "success": True,
            "message": "Índice recarregado com sucesso",
            **version_compare_service.get_index_metadata(),
        }
    except Exception as e:
        logger.error(f"Error reloading version compare index: {e}")
        raise fastapi.HTTPException(status_code=500, detail=str(e))


@api_router.post("/version-compare")
async def version_compare(log_file: fastapi.UploadFile = fastapi.File(...)):
    """Compara versoes do extrato do cliente contra o repositorio corporativo."""
    try:
        request_start = monotonic()
        file_content = await log_file.read()
        decode_start = monotonic()
        content = decode_uploaded_text(file_content)
        decode_ms = round((monotonic() - decode_start) * 1000, 2)
        compare_start = monotonic()
        result = version_compare_service.compare_content(content)
        compare_ms = round((monotonic() - compare_start) * 1000, 2)
        existing_timings = result.get("timings") or {}
        result["timings"] = {
            **existing_timings,
            "read_file_bytes": len(file_content),
            "decode_uploaded_text_ms": decode_ms,
            "endpoint_compare_ms": compare_ms,
            "total_request_ms": round((monotonic() - request_start) * 1000, 2),
        }
        return {"success": True, **result}
    except ValueError as e:
        raise fastapi.HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Error running version compare: {e}")
        raise fastapi.HTTPException(status_code=500, detail=str(e))

def extract_timestamp_from_line(line: str) -> str:
    """Extrai timestamp de uma linha de log."""
    timestamp_patterns = [
        r'(\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2})',
        r'(\d{2}/\d{2}/\d{4} \d{2}:\d{2}:\d{2})',
        r'(\d{4}/\d{2}/\d{2} \d{2}:\d{2}:\d{2})',
        r'(\w{3} \d{2} \d{2}:\d{2}:\d{2})',
    ]
    
    for pattern in timestamp_patterns:
        match = re.search(pattern, line)
        if match:
            return match.group(1)
    return None

# =====================================================================
# Issue Control & Evidence Register (from aberturadetickets)
# =====================================================================

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)


class TicketData(BaseModel):
    ticket_number: str
    situation: str
    issue: Optional[str] = None
    client_name: str
    client_version: str
    database_type: str
    routine_program: str
    occurrence_type: str
    simulated_internally: bool
    simulation_base: Optional[str] = None
    occurrence_description: str
    expected_result: str
    program_needed: bool
    notification_emails: Optional[str] = None


class IssueData(BaseModel):
    id: Optional[str] = Field(default_factory=lambda: str(uuid.uuid4()))
    data_criacao: Optional[str] = Field(default_factory=lambda: datetime.now().strftime('%d/%m/%Y'))
    ticket: str
    issue: str
    cliente: str
    rotina: str
    situacao: str
    status: str
    liberado_versoes: Optional[str] = None


class IssueCreate(BaseModel):
    ticket: str
    issue: str
    cliente: str
    rotina: str
    situacao: str
    status: str
    liberado_versoes: Optional[str] = None


class IssueUpdate(BaseModel):
    ticket: Optional[str] = None
    issue: Optional[str] = None
    cliente: Optional[str] = None
    rotina: Optional[str] = None
    situacao: Optional[str] = None
    status: Optional[str] = None
    liberado_versoes: Optional[str] = None


def sanitize_filename_component(value: str, fallback: str) -> str:
    sanitized = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '_', value or '')
    sanitized = re.sub(r'\s+', ' ', sanitized).strip().strip('.')
    return sanitized[:120] or fallback


def plain_text_to_reportlab(text: str) -> str:
    if not text or text.strip() == '':
        return ''
    content = escape(html.unescape(text)).replace('\r\n', '\n').replace('\r', '\n')
    return content.replace('\n', '<br/>')


def sanitize_html_for_doc(html_content: str) -> str:
    if not html_content or html_content.strip() == '':
        return '<p>Descricao nao fornecida.</p>'
    content = html.unescape(html_content)
    content = re.sub(r'<(script|style)\b[^>]*>.*?</\1>', '', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'<img\b[^>]*>', '', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'\sstyle\s*=\s*"[^"]*"', '', content, flags=re.IGNORECASE)
    content = re.sub(r"\sstyle\s*=\s*'[^']*'", '', content, flags=re.IGNORECASE)
    content = re.sub(r'<(?!/?(?:p|br|strong|b|em|i|u|ul|ol|li|div)\b)[^>]+>', '', content, flags=re.IGNORECASE)
    content = content.strip()
    return content or '<p>Descricao nao fornecida.</p>'


def html_to_reportlab_ticket(html_content: str) -> str:
    if not html_content or html_content.strip() == '':
        return ""
    content = html.unescape(html_content).replace('\r\n', '\n').replace('\r', '\n')
    markers = {
        'line_break': '__LINE_BREAK__',
        'bold_open': '__BOLD_OPEN__', 'bold_close': '__BOLD_CLOSE__',
        'italic_open': '__ITALIC_OPEN__', 'italic_close': '__ITALIC_CLOSE__',
        'underline_open': '__UNDERLINE_OPEN__', 'underline_close': '__UNDERLINE_CLOSE__',
    }
    content = re.sub(r'<img\b[^>]*>', '', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'<br\s*/?>', markers['line_break'], content, flags=re.IGNORECASE)
    content = re.sub(r'</?(div|p|ul|ol|table|tr|blockquote)\b[^>]*>', markers['line_break'], content, flags=re.IGNORECASE)
    content = re.sub(r'<li\b[^>]*>', f"{markers['line_break']}* ", content, flags=re.IGNORECASE)
    content = re.sub(r'</li>', '', content, flags=re.IGNORECASE)
    for pat, rep in [
        (r'<(strong|b)\b[^>]*>', markers['bold_open']), (r'</(strong|b)>', markers['bold_close']),
        (r'<(em|i)\b[^>]*>', markers['italic_open']), (r'</(em|i)>', markers['italic_close']),
        (r'<u\b[^>]*>', markers['underline_open']), (r'</u>', markers['underline_close']),
    ]:
        content = re.sub(pat, rep, content, flags=re.IGNORECASE)
    content = re.sub(r'<[^>]+>', '', content)
    content = re.sub(r'\n+', markers['line_break'], content)
    content = re.sub(r'[ \t\f\v]+', ' ', content)
    content = re.sub(fr'\s*{markers["line_break"]}\s*', markers['line_break'], content)
    content = re.sub(fr'({markers["line_break"]}){{3,}}', markers['line_break'] * 2, content)
    content = content.strip()
    if not content:
        return ""
    content = escape(content)
    for marker, replacement in {
        markers['bold_open']: '<b>', markers['bold_close']: '</b>',
        markers['italic_open']: '<i>', markers['italic_close']: '</i>',
        markers['underline_open']: '<u>', markers['underline_close']: '</u>',
        markers['line_break']: '<br/>',
    }.items():
        content = content.replace(marker, replacement)
    content = re.sub(r'(<br/>){3,}', '<br/><br/>', content)
    content = re.sub(r'^(<br/>)+', '', content)
    content = re.sub(r'(<br/>)+$', '', content)
    return content


def _add_docx_run(paragraph, text, bold=False, italic=False, underline=False):
    if not text:
        return
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)


def _append_docx_inline(paragraph, node, bold=False, italic=False, underline=False):
    if isinstance(node, str):
        _add_docx_run(paragraph, node, bold=bold, italic=italic, underline=underline)
        return
    tag = (node.tag or '').lower() if hasattr(node, 'tag') else ''
    nb = bold or tag in {'strong', 'b'}
    ni = italic or tag in {'em', 'i'}
    nu = underline or tag == 'u'
    if getattr(node, 'text', None):
        _add_docx_run(paragraph, node.text, bold=nb, italic=ni, underline=nu)
    for child in node:
        ct = (child.tag or '').lower() if hasattr(child, 'tag') else ''
        if ct == 'br':
            paragraph.add_run().add_break()
        else:
            _append_docx_inline(paragraph, child, bold=nb, italic=ni, underline=nu)
        if child.tail:
            _add_docx_run(paragraph, child.tail, bold=nb, italic=ni, underline=nu)


def _append_html_to_docx(container, html_content):
    sanitized = sanitize_html_for_doc(html_content)
    try:
        fragments = lxml_html.fragments_fromstring(sanitized)
    except Exception:
        fragments = [html.unescape(re.sub(r'<[^>]+>', '', sanitized))]
    has_content = False
    def add_p(element, style=None):
        nonlocal has_content
        p = container.add_paragraph(style=style)
        _append_docx_inline(p, element)
        if not p.text.strip():
            p.add_run(' ')
        has_content = True
    for fragment in fragments:
        if isinstance(fragment, str):
            t = fragment.strip()
            if t:
                p = container.add_paragraph()
                _add_docx_run(p, t)
                has_content = True
            continue
        tag = (fragment.tag or '').lower()
        if tag in {'p', 'div'}:
            add_p(fragment)
        elif tag in {'ul', 'ol'}:
            ls = 'List Bullet' if tag == 'ul' else 'List Number'
            for item in fragment:
                if (getattr(item, 'tag', '') or '').lower() == 'li':
                    add_p(item, style=ls)
        elif tag == 'li':
            add_p(fragment, style='List Bullet')
        else:
            add_p(fragment)
    if not has_content:
        p = container.add_paragraph()
        _add_docx_run(p, 'Descricao nao fornecida.')


def _set_docx_cell_text(cell, label, value):
    cell.text = ''
    p = cell.paragraphs[0]
    if not label:
        p.add_run(' ')
        return
    lr = p.add_run(f'{label}: ')
    lr.bold = True
    lr.font.name = 'Arial'
    lr.font.size = Pt(10.5)
    vr = p.add_run((value or '').strip() or 'Nao informado')
    vr.font.name = 'Arial'
    vr.font.size = Pt(10.5)


def build_docx_document(ticket_data: TicketData, output_path: Path):
    document = DocxDocument()
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run('TOTVS')
    tr.bold = True
    tr.font.name = 'Arial'
    tr.font.size = Pt(24)
    tr.font.color.rgb = RGBColor(15, 42, 76)
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('Portal de Registro de Evidencias')
    sr.font.name = 'Arial'
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(41, 68, 111)
    document.add_paragraph()

    def add_heading(text):
        p = document.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(15, 42, 76)

    def add_info_table(rows):
        table = document.add_table(rows=len(rows), cols=2)
        table.style = 'Table Grid'
        table.autofit = True
        for ri, row in enumerate(rows):
            for ci in range(2):
                lbl, val = row[ci] if ci < len(row) else ('', '')
                _set_docx_cell_text(table.rows[ri].cells[ci], lbl, val)
        document.add_paragraph()

    add_heading('Informacoes do Ticket')
    add_info_table([
        [('No do Ticket', ticket_data.ticket_number), ('Cliente', ticket_data.client_name)],
        [('Titulo do Chamado', ticket_data.situation), ('Issue', ticket_data.issue)],
        [('Emails para Notificacao', ticket_data.notification_emails), ('', '')],
    ])
    add_heading('Configuracao Tecnica')
    add_info_table([
        [('Rotina/Programa', ticket_data.routine_program), ('Tipo de Ocorrencia', ticket_data.occurrence_type)],
        [('Versao', ticket_data.client_version), ('Banco de Dados', ticket_data.database_type)],
    ])
    add_heading('Analise Tecnica')
    add_info_table([
        [('Situacao simulada internamente?', 'Sim' if ticket_data.simulated_internally else 'Nao'),
         ('Necessario Programa de Acerto?', 'Sim' if ticket_data.program_needed else 'Nao')],
        [('Base de Simulacao', ticket_data.simulation_base), ('', '')],
    ])
    add_heading('Descricao Detalhada')
    _append_html_to_docx(document, ticket_data.occurrence_description)
    document.add_paragraph()
    add_heading('Resultado Esperado')
    ep = document.add_paragraph()
    et = (ticket_data.expected_result or '').strip() or 'Resultado esperado nao fornecido.'
    for idx, line in enumerate(et.replace('\r\n', '\n').replace('\r', '\n').split('\n')):
        if idx > 0:
            ep.add_run().add_break()
        _add_docx_run(ep, line)
    document.save(str(output_path))


@api_router.post("/upload-files")
async def upload_files(files: List[fastapi.UploadFile] = fastapi.File(...), session_id: Optional[str] = fastapi.Form(None)):
    try:
        uploaded = []
        session_id = session_id if session_id and (UPLOAD_DIR / session_id).exists() else str(uuid.uuid4())
        session_dir = UPLOAD_DIR / session_id
        session_dir.mkdir(exist_ok=True)
        for f in files:
            contents = await f.read()
            if len(contents) > 20 * 1024 * 1024:
                raise fastapi.HTTPException(status_code=413, detail=f"Arquivo {f.filename} excede o limite de 20MB")
            fp = session_dir / f.filename
            async with aiofiles.open(fp, 'wb') as out:
                await out.write(contents)
            uploaded.append({"filename": f.filename, "path": str(fp), "size": len(contents)})
        return {"session_id": session_id, "files": uploaded, "message": f"Upload de {len(files)} arquivo(s) realizado com sucesso"}
    except fastapi.HTTPException:
        raise
    except Exception as e:
        raise fastapi.HTTPException(status_code=500, detail=f"Falha no upload: {str(e)}")


@api_router.delete("/uploaded-files/{session_id}")
async def delete_uploaded_file(session_id: str, filename: str):
    try:
        session_dir = UPLOAD_DIR / session_id
        fp = session_dir / Path(filename).name
        if not session_dir.exists() or not fp.exists():
            raise fastapi.HTTPException(status_code=404, detail="Arquivo nao encontrado")
        fp.unlink()
        if not any(session_dir.iterdir()):
            session_dir.rmdir()
        return {"message": "Arquivo removido com sucesso"}
    except fastapi.HTTPException:
        raise
    except Exception as e:
        raise fastapi.HTTPException(status_code=500, detail=str(e))


@api_router.post("/generate-pdf")
async def generate_pdf(
    ticket_number: str = fastapi.Form(...),
    situation: str = fastapi.Form(...),
    issue: Optional[str] = fastapi.Form(None),
    client_name: str = fastapi.Form(...),
    client_version: str = fastapi.Form(...),
    database_type: str = fastapi.Form(...),
    routine_program: str = fastapi.Form(...),
    occurrence_type: str = fastapi.Form(...),
    simulated_internally: bool = fastapi.Form(...),
    simulation_base: Optional[str] = fastapi.Form(None),
    occurrence_description: str = fastapi.Form(...),
    expected_result: str = fastapi.Form(...),
    program_needed: bool = fastapi.Form(...),
    notification_emails: Optional[str] = fastapi.Form(None),
    session_id: Optional[str] = fastapi.Form(None),
):
    try:
        temp_dir = Path(tempfile.mkdtemp())
        ticket_data = TicketData(
            ticket_number=ticket_number, situation=situation, issue=issue,
            client_name=client_name, client_version=client_version,
            database_type=database_type, routine_program=routine_program,
            occurrence_type=occurrence_type, simulated_internally=simulated_internally,
            simulation_base=simulation_base, occurrence_description=occurrence_description,
            expected_result=expected_result, program_needed=program_needed,
            notification_emails=notification_emails,
        )
        safe_name = f"{sanitize_filename_component(ticket_data.ticket_number, 'ticket')}_{sanitize_filename_component(ticket_data.situation, 'evidencias')}"
        pdf_filename = f"{safe_name}.pdf"
        pdf_path = temp_dir / pdf_filename
        doc_filename = f"{safe_name}.docx"
        doc_path = temp_dir / doc_filename

        # Build PDF
        doc = BaseDocTemplate(str(pdf_path), pagesize=A4)
        frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height - 120,
                      leftPadding=0, rightPadding=0, topPadding=20, bottomPadding=0)
        def draw_header(canvas, doc):
            canvas.setFillColor(colors.black)
            canvas.rect(0, A4[1] - 100, A4[0], 100, fill=1, stroke=0)
            canvas.setFillColor(colors.white)
            canvas.setFont("Helvetica-Bold", 24)
            canvas.drawCentredString(A4[0] / 2, A4[1] - 58, "TOTVS")
            canvas.setFillColor(colors.HexColor('#dbeafe'))
            canvas.setFont("Helvetica", 10)
            canvas.drawCentredString(A4[0] / 2, A4[1] - 74, "Portal de Registro de Evidencias")
        tmpl = PageTemplate(id='main', frames=[frame], onPage=draw_header)
        doc.addPageTemplates([tmpl])
        styles = getSampleStyleSheet()
        section_title_style = ParagraphStyle('SectionTitle', parent=styles['Heading2'], fontSize=16, spaceBefore=25, spaceAfter=15, textColor=colors.HexColor('#0f2a4c'), fontName='Helvetica-Bold')
        field_style = ParagraphStyle('FieldStyle', parent=styles['Normal'], fontSize=11, spaceAfter=8, fontName='Helvetica')
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24, spaceAfter=5, textColor=colors.HexColor('#0f2a4c'), alignment=1, fontName='Helvetica-Bold')
        subtitle_style = ParagraphStyle('CustomSubtitle', parent=styles['Normal'], fontSize=14, spaceAfter=30, textColor=colors.black, alignment=1, fontName='Helvetica')
        story = [Spacer(1, 20)]
        story.append(Paragraph("Portal de Registro de Evidencias", title_style))
        story.append(Paragraph("Sistema de Controle de Tickets para Desenvolvimento", subtitle_style))

        def section_table(title, pairs):
            story.append(Paragraph(title, section_title_style))
            rows = []
            for i in range(0, len(pairs), 2):
                row = []
                if i < len(pairs):
                    l, v = pairs[i]
                    row.append(Paragraph(f"<b>{l}:</b><br/>{v}", field_style))
                else:
                    row.append("")
                if i + 1 < len(pairs):
                    l, v = pairs[i + 1]
                    row.append(Paragraph(f"<b>{l}:</b><br/>{v}", field_style))
                else:
                    row.append("")
                rows.append(row)
            if rows:
                t = Table(rows, colWidths=[doc.width / 2 - 10, doc.width / 2 - 10])
                t.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 5),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 15),
                ]))
                story.append(t)

        ti = [("No do Ticket", ticket_data.ticket_number), ("Cliente", ticket_data.client_name), ("Titulo", ticket_data.situation)]
        if ticket_data.issue:
            ti.append(("Issue", ticket_data.issue))
        if ticket_data.notification_emails:
            ti.append(("Emails", ticket_data.notification_emails))
        section_table("Informacoes do Ticket", ti)
        section_table("Configuracao Tecnica", [
            ("Rotina/Programa", ticket_data.routine_program), ("Tipo de Ocorrencia", ticket_data.occurrence_type),
            ("Versao", ticket_data.client_version), ("Banco de Dados", ticket_data.database_type),
        ])
        af = [("Simulada internamente?", "Sim" if ticket_data.simulated_internally else "Nao"),
              ("Programa de acerto?", "Sim" if ticket_data.program_needed else "Nao")]
        if ticket_data.simulated_internally and ticket_data.simulation_base:
            af.insert(1, ("Base de Simulacao", ticket_data.simulation_base))
        section_table("Analise Tecnica", af)

        story.append(Paragraph("Descricao Detalhada", section_title_style))
        story.append(Paragraph("<b>Situacao:</b>", field_style))
        try:
            pd = html_to_reportlab_ticket(ticket_data.occurrence_description) or "Descricao nao fornecida."
            story.append(Paragraph(pd, field_style))
        except Exception:
            story.append(Paragraph("Erro ao processar descricao.", field_style))
        story.append(Spacer(1, 10))
        story.append(Paragraph("<b>Resultado Esperado:</b>", field_style))
        try:
            er = plain_text_to_reportlab(ticket_data.expected_result) or "Resultado esperado nao fornecido."
            story.append(Paragraph(er, field_style))
        except Exception:
            story.append(Paragraph("Erro ao processar resultado esperado.", field_style))
        doc.build(story)

        # Build DOCX
        build_docx_document(ticket_data, doc_path)

        # Save to issues collection
        try:
            uploaded_file_names = []
            if session_id:
                sd = UPLOAD_DIR / session_id
                if sd.exists():
                    uploaded_file_names = [p.name for p in sd.iterdir() if p.is_file()]
            existing = await find_issue_record({"ticket": ticket_data.ticket_number})
            issue_doc = {
                "ticket": ticket_data.ticket_number,
                "issue": ticket_data.issue or "Nao informado",
                "cliente": ticket_data.client_name,
                "rotina": ticket_data.routine_program,
                "situacao": ticket_data.situation,
                "status": "Aberto",
                "liberado_versoes": "",
                "data_criacao": datetime.now().strftime('%d/%m/%Y'),
                "portal_detalhes": {
                    "ticket_number": ticket_data.ticket_number,
                    "situation": ticket_data.situation,
                    "issue": ticket_data.issue,
                    "client_name": ticket_data.client_name,
                    "client_version": ticket_data.client_version,
                    "database_type": ticket_data.database_type,
                    "routine_program": ticket_data.routine_program,
                    "occurrence_type": ticket_data.occurrence_type,
                    "simulated_internally": ticket_data.simulated_internally,
                    "simulation_base": ticket_data.simulation_base,
                    "occurrence_description": ticket_data.occurrence_description,
                    "expected_result": ticket_data.expected_result,
                    "program_needed": ticket_data.program_needed,
                    "notification_emails": ticket_data.notification_emails,
                    "arquivos_anexados": uploaded_file_names,
                },
            }
            if existing:
                await update_issue_records({"ticket": ticket_data.ticket_number}, issue_doc)
            else:
                issue_doc["id"] = str(uuid.uuid4())
                await save_issue_record(issue_doc)
        except Exception as save_err:
            logger.error(f"Failed to save issue: {save_err}")

        # Create ZIP package
        zip_path = temp_dir / f"{safe_name}.zip"
        with zipfile.ZipFile(zip_path, 'w') as zf:
            zf.write(pdf_path, pdf_filename)
            zf.write(doc_path, doc_filename)
            if session_id:
                sd = UPLOAD_DIR / session_id
                if sd.exists():
                    for fp in sd.iterdir():
                        zf.write(fp, fp.name)
        return FileResponse(path=str(zip_path), filename=f"{safe_name}.zip", media_type="application/zip")
    except fastapi.HTTPException:
        raise
    except Exception as e:
        raise fastapi.HTTPException(status_code=500, detail=f"Falha na geracao de PDF: {str(e)}")


@api_router.delete("/cleanup-session/{session_id}")
async def cleanup_session(session_id: str):
    try:
        sd = UPLOAD_DIR / session_id
        if sd.exists():
            shutil.rmtree(sd)
        return {"message": "Sessao limpa com sucesso"}
    except Exception as e:
        raise fastapi.HTTPException(status_code=500, detail=str(e))


@api_router.get("/issues")
async def get_issues():
    try:
        issues = await list_issue_records()
        result = []
        for iss in issues:
            result.append({
                "id": iss.get("id", ""),
                "data_criacao": iss.get("data_criacao", ""),
                "ticket": iss.get("ticket", ""),
                "issue": iss.get("issue", ""),
                "cliente": iss.get("cliente", ""),
                "rotina": iss.get("rotina", ""),
                "situacao": iss.get("situacao", ""),
                "status": iss.get("status", ""),
                "liberado_versoes": iss.get("liberado_versoes", ""),
            })
        return result
    except Exception as e:
        raise fastapi.HTTPException(status_code=500, detail=str(e))


@api_router.post("/issues")
async def create_issue(issue: IssueCreate):
    try:
        issue_data = IssueData(**issue.dict())
        doc = issue_data.dict()
        await save_issue_record(doc)
        doc.pop("_id", None)
        return doc
    except Exception as e:
        raise fastapi.HTTPException(status_code=500, detail=str(e))


@api_router.put("/issues/{issue_id}")
async def update_issue(issue_id: str, update_data: IssueUpdate):
    try:
        ud = {k: v for k, v in update_data.dict().items() if v is not None}
        if not ud:
            raise fastapi.HTTPException(status_code=400, detail="Nenhum dado para atualizar")
        matched_count = await update_issue_records({"id": issue_id}, ud)
        if matched_count == 0:
            raise fastapi.HTTPException(status_code=404, detail="Issue nao encontrada")
        return {"message": "Issue atualizada com sucesso"}
    except fastapi.HTTPException:
        raise
    except Exception as e:
        raise fastapi.HTTPException(status_code=500, detail=str(e))


@api_router.delete("/issues/{issue_id}")
async def delete_issue(issue_id: str):
    try:
        deleted_count = await delete_issue_records({"id": issue_id})
        if deleted_count == 0:
            raise fastapi.HTTPException(status_code=404, detail="Issue nao encontrada")
        return {"message": "Issue removida com sucesso"}
    except fastapi.HTTPException:
        raise
    except Exception as e:
        raise fastapi.HTTPException(status_code=500, detail=str(e))


@api_router.post("/import-csv")
async def import_csv(file: fastapi.UploadFile = fastapi.File(...)):
    try:
        if not file.filename.lower().endswith('.csv'):
            raise fastapi.HTTPException(status_code=400, detail="Arquivo deve ser CSV")
        contents = await file.read()
        csv_content = contents.decode('utf-8')
        reader = csv.DictReader(io.StringIO(csv_content), delimiter=';')
        expected = {'Data', 'Ticket', 'Issue', 'Cliente', 'Rotina', 'Situacao', 'Status', 'Liberado_Versoes'}
        if not reader.fieldnames or set(reader.fieldnames) != expected:
            raise fastapi.HTTPException(status_code=400, detail=f"CSV deve ter headers: {';'.join(sorted(expected))}")
        imported_count = 0
        errors = []
        for rn, row in enumerate(reader, start=2):
            try:
                if not row.get('Ticket') or not row.get('Issue'):
                    errors.append(f"Linha {rn}: Ticket e Issue obrigatorios")
                    continue
                existing = await find_issue_record({"ticket": row['Ticket'].strip()})
                if existing:
                    errors.append(f"Linha {rn}: Ticket {row['Ticket'].strip()} ja existe")
                    continue
                doc = IssueData(
                    data_criacao=row.get('Data', '').strip() or datetime.now().strftime('%d/%m/%Y'),
                    ticket=row['Ticket'].strip(),
                    issue=row['Issue'].strip(),
                    cliente=(row.get('Cliente') or '').strip(),
                    rotina=(row.get('Rotina') or '').strip(),
                    situacao=(row.get('Situacao') or '').strip(),
                    status=(row.get('Status') or 'Aberto').strip(),
                    liberado_versoes=(row.get('Liberado_Versoes') or '').strip(),
                ).dict()
                await save_issue_record(doc)
                imported_count += 1
            except Exception as ex:
                errors.append(f"Linha {rn}: {str(ex)}")
        resp = {"imported_count": imported_count}
        if errors:
            resp["errors"] = errors[:10]
            resp["total_errors"] = len(errors)
        return resp
    except fastapi.HTTPException:
        raise
    except Exception as e:
        raise fastapi.HTTPException(status_code=500, detail=str(e))


@api_router.get("/export-csv")
async def export_csv():
    try:
        issues = await list_issue_records()
        output = io.StringIO()
        writer = csv.writer(output, delimiter=';')
        writer.writerow(['Data', 'Ticket', 'Issue', 'Cliente', 'Rotina', 'Situacao', 'Status', 'Liberado_Versoes'])
        for iss in issues:
            writer.writerow([
                iss.get('data_criacao', ''), iss.get('ticket', ''), iss.get('issue', ''),
                iss.get('cliente', ''), iss.get('rotina', ''), iss.get('situacao', ''),
                iss.get('status', ''), iss.get('liberado_versoes', ''),
            ])
        csv_content = output.getvalue()
        output.close()
        return Response(content=csv_content.encode('utf-8'), media_type='text/csv',
                        headers={"Content-Disposition": "attachment; filename=issues.csv"})
    except Exception as e:
        raise fastapi.HTTPException(status_code=500, detail=str(e))


# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
